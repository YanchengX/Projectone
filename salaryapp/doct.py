from docx import Document
from docx.shared import RGBColor
from docx.shared import Cm, Pt
class Preview:
    def __init__(self, empdata) -> None:
        self.empdata = empdata
        data = self.empdata.fetchone()
        print(data)
    # simulate doc      
        doc = Document()
        #section
        section = doc.sections[0]
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)

        #normal
        doc.add_heading('員工薪資明細表', 0)
        sdate = doc.add_paragraph('月份：')
        sdate.add_run('111')
        sdate.add_run('年')
        sdate.add_run('12')
        sdate.add_run('月')
        table_normal = doc.add_table(rows=17, cols=4)
        table_normal.style = 'Table Grid'

        merge_a = table_normal.cell(0, 0)
        merge_b = table_normal.cell(0, 1)
        merge_f = merge_a.merge(merge_b)
        merge_a = table_normal.cell(1, 0)
        merge_b = table_normal.cell(1, 1)
        merge_s = merge_a.merge(merge_b)
        merge_a = table_normal.cell(2, 0)
        merge_b = table_normal.cell(8, 0)
        merge_0 = merge_a.merge(merge_b)
        merge_a = table_normal.cell(9, 0)
        merge_b = table_normal.cell(15, 0)
        merge_1 = merge_a.merge(merge_b)
        merge_a = table_normal.cell(16, 0)
        merge_b = table_normal.cell(16, 1)
        merge_2 = merge_a.merge(merge_b)

        table_normal.cell(0,0).text = '編   號'
        table_normal.cell(0,2).text = 'A01'
        table_normal.cell(1,0).text = '姓   名'
        table_normal.cell(1,2).text = 'yan1100923'
        table_normal.cell(2,4).text = '應領薪資金額'
        table_normal.cell(2,1).text = '基本薪資(月薪)'
        table_normal.cell(2,2).text = '1250'
        table_normal.cell(2,3).text = '23650'

        table_normal.cell(3,1).text = '伙食津貼(60)'
        table_normal.cell(3,2).text = '60*21天'
        table_normal.cell(3,3).text = '1260'

        table_normal.cell(4,1).text = '全勤獎金(1000)'
        table_normal.cell(4,2).text = '0天'
        table_normal.cell(4,3).text = '0'

        table_normal.cell(5,1).text = '開門津貼'
        table_normal.cell(5,2).text = '0天'
        table_normal.cell(5,3).text = '0'

        table_normal.cell(6,1).text = '責任津貼'
        table_normal.cell(6,2).text = '0天'
        table_normal.cell(6,3).text = '0'

        table_normal.cell(7,1).text = '其他'
        table_normal.cell(7,2).text = '0天'
        table_normal.cell(7,3).text = '0'

        table_normal.cell(8,1).text = '小記'
        table_normal.cell(8,2).text = '0天'
        table_normal.cell(8,3).text = '0'

        table_normal.cell(9,1).text = '勞保費'
        table_normal.cell(9,2).text = '0天'
        table_normal.cell(9,3).text = '0'
        table_normal.cell(9,4).text = '應扣金額'

        table_normal.cell(10,1).text = '健保費'
        table_normal.cell(10,2).text = '0天'
        table_normal.cell(10,3).text = '0'

        table_normal.cell(11,1).text = '請假'
        table_normal.cell(11,2).text = '0天'
        table_normal.cell(11,3).text = '0'

        table_normal.cell(12,1).text = '介之'
        table_normal.cell(12,2).text = '0天'
        table_normal.cell(12,3).text = '0'

        table_normal.cell(13,1).text = '帶定伙食費'
        table_normal.cell(13,2).text = '0天'
        table_normal.cell(13,3).text = '0'

        table_normal.cell(14,1).text = '其他'
        table_normal.cell(14,2).text = '0天'
        table_normal.cell(14,3).text = '0'

        table_normal.cell(15,1).text = '小記'
        table_normal.cell(15,2).text = '0天'
        table_normal.cell(15,3).text = '0'

        table_normal.cell(16,1).text = '總計'
        table_normal.cell(16,2).text = '0天'
        table_normal.cell(16,3).text = '0'

        #overtime
        doc.add_heading('員工月薪加班薪資明細表',0)
        sdateo = doc.add_paragraph('月份：')
        sdateo.add_run('111')
        sdateo.add_run('年')
        sdateo.add_run('12')
        sdateo.add_run('月')
        table_overtime = doc.add_table(rows=19, cols=4)
        table_overtime.style = 'Table Grid'

        a = table_overtime.cell(0, 0)
        b = table_overtime.cell(0, 1)
        merge0 = a.merge(b)
        a = table_overtime.cell(1, 0)
        b = table_overtime.cell(1, 1)
        merge1 = a.merge(b)
        a = table_overtime.cell(2, 0)
        b = table_overtime.cell(12, 0)
        merge2 = a.merge(b)

        a = table_overtime.cell(13, 0)
        b = table_overtime.cell(13, 1)
        merge3 = a.merge(b)
        a = table_overtime.cell(14, 0)
        b = table_overtime.cell(14, 1)
        merge4 = a.merge(b)
        a = table_overtime.cell(15, 0)
        b = table_overtime.cell(15, 1)
        merge5 = a.merge(b)
        a = table_overtime.cell(16, 0)
        b = table_overtime.cell(16, 1)
        merge6 = a.merge(b)
        a = table_overtime.cell(17, 0)
        b = table_overtime.cell(17, 1)
        merge7 = a.merge(b)
        a = table_overtime.cell(18, 0)
        b = table_overtime.cell(18, 1)
        merge8 = a.merge(b)

        table_overtime.cell(0,1).text = '編號'
        table_overtime.cell(0,2).text = 'A02'
        table_overtime.cell(1,1).text = '姓名'
        table_overtime.cell(1,2).text = '彥董1100912'

        table_overtime.cell(2,1).text = '平日加班前兩小時(1.34)'
        table_overtime.cell(2,2).text = '176x15小時'
        table_overtime.cell(2,3).text = '1974'
        table_overtime.cell(2,4).text = '應領加班薪資金額'

        table_overtime.cell(3,1).text = '平日加班後兩小時(1.67)'
        table_overtime.cell(3,2).text = '176x15小時'
        table_overtime.cell(3,3).text = '0'

        table_overtime.cell(4,1).text = '平日加班伙食津貼'
        table_overtime.cell(4,2).text = '176x15小時'
        table_overtime.cell(4,3).text = '0'

        table_overtime.cell(5,1).text = '休息日加班(六)(前1.34後1.67)'
        table_overtime.cell(5,2).text = '176x15小時'
        table_overtime.cell(5,3).text = '0'

        table_overtime.cell(6,1).text = '休息日加班(六)伙食津貼'
        table_overtime.cell(6,2).text = '176x15小時'
        table_overtime.cell(6,3).text = '0'

        table_overtime.cell(7,1).text = '例假日加班(日)(一天+補修一天)'
        table_overtime.cell(7,2).text = '176x15小時'
        table_overtime.cell(7,3).text = '0'

        table_overtime.cell(8,1).text = '例假日加班(日)伙食津貼'
        table_overtime.cell(8,2).text = '176x15小時'
        table_overtime.cell(8,3).text = '0'

        table_overtime.cell(9,1).text = '國定假日加班'
        table_overtime.cell(9,2).text = '176x15小時'
        table_overtime.cell(9,3).text = '0'

        table_overtime.cell(10,1).text = '國定假日加班伙食津貼'
        table_overtime.cell(10,2).text = '176x15小時'
        table_overtime.cell(10,3).text = '0'

        table_overtime.cell(11,1).text = '其他'
        table_overtime.cell(11,2).text = '176x15小時'
        table_overtime.cell(11,3).text = '0'

        table_overtime.cell(12,1).text = '小記'
        table_overtime.cell(12,2).text = '176x15小時'
        table_overtime.cell(12,3).text = '0'

        table_overtime.cell(13,1).text = '總計'
        table_overtime.cell(13,2).text = '176x15小時'
        table_overtime.cell(13,3).text = '0'

        table_overtime.cell(14,1).text = '總領金額'
        table_overtime.cell(14,2).text = '176x15小時'
        table_overtime.cell(14,3).text = '0'

        table_overtime.cell(15,1).text = '轉帳'
        table_overtime.cell(15,2).text = '176x15小時'
        table_overtime.cell(15,3).text = '0'

        table_overtime.cell(16,1).text = '現金'
        table_overtime.cell(16,2).text = '176x15小時'
        table_overtime.cell(16,3).text = '0'

        table_overtime.cell(17,1).text = '勞退6%'
        table_overtime.cell(17,2).text = '176x15小時'
        table_overtime.cell(17,3).text = '0'

        table_overtime.cell(18,1).text = '特休'
        table_overtime.cell(18,2).text = '176x15小時'
        table_overtime.cell(18,3).text = '0'

        doc.save('doct.docx')





# # read doc
# doc = Document('doct.docx')
# for paragraph in doc.paragraphs:
#     print(paragraph.text)
# for i in range(0,len(doc.paragraphs[3].runs)):
#     print(doc.paragraphs[3].runs[i].text)


#create doc
# document = Document()

# document.add_heading("Testing doc",0)

# table = document.add_table(rows = 5, cols = 2, style = 'Table Grid')
# table.rows[0].cells[0].text = 'fuck tou'
# document.save('doct.docx')